"""
HireFlow AI - Resume Parser Service
Extracts text from PDF and DOCX files for AI analysis
"""

import os
from flask import current_app
from PyPDF2 import PdfReader
from docx import Document


class ResumeParser:
    """Service for parsing resume files"""

    @staticmethod
    def extract_text(file_path):
        """
        Extract text from resume file
        Supports PDF and DOCX formats
        """
        if not file_path or not os.path.exists(file_path):
            return {'error': 'File not found', 'success': False}
        
        ext = file_path.rsplit('.', 1)[1].lower() if '.' in file_path else ''
        
        try:
            if ext == 'pdf':
                return ResumeParser._parse_pdf(file_path)
            elif ext in ('docx', 'doc'):
                return ResumeParser._parse_docx(file_path)
            else:
                return {'error': f'Unsupported file format: {ext}', 'success': False}
        except Exception as e:
            current_app.logger.error(f'Error parsing resume: {str(e)}')
            return {'error': f'Failed to parse file: {str(e)}', 'success': False}

    @staticmethod
    def _parse_pdf(file_path):
        """Extract text from PDF file"""
        text_content = []
        metadata = {}
        
        with open(file_path, 'rb') as file:
            pdf = PdfReader(file)
            metadata = {
                'pages': len(pdf.pages),
                'pdf_version': pdf.pdf_header if hasattr(pdf, 'pdf_header') else None,
            }
            
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        
        full_text = '\n\n'.join(text_content)
        
        return {
            'text': full_text,
            'metadata': metadata,
            'success': True,
            'format': 'pdf'
        }

    @staticmethod
    def _parse_docx(file_path):
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    text_content.append(' | '.join(row_data))
        
        full_text = '\n'.join(text_content)
        
        return {
            'text': full_text,
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            },
            'success': True,
            'format': 'docx'
        }

    @staticmethod
    def extract_sections(text):
        """Extract common resume sections from text"""
        if not text:
            return {}
        
        sections = {
            'contact_info': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'projects': '',
            'certifications': '',
            'languages': ''
        }
        
        lines = text.split('\n')
        current_section = 'contact_info'
        
        section_keywords = {
            'summary': ['professional summary', 'profile', 'about me', 'career objective'],
            'experience': ['experience', 'work experience', 'employment', 'work history'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'core competencies', 'expertise'],
            'projects': ['projects', 'project experience', 'key projects'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'languages': ['languages', 'language proficiency']
        }
        
        for line in lines:
            line_lower = line.lower().strip()
            found_section = False
            
            for section, keywords in section_keywords.items():
                for keyword in keywords:
                    if keyword in line_lower and len(line) < 100:
                        current_section = section
                        found_section = True
                        break
                if found_section:
                    break
            
            if current_section in sections:
                if sections[current_section]:
                    sections[current_section] += '\n' + line
                else:
                    sections[current_section] = line
        
        return sections
